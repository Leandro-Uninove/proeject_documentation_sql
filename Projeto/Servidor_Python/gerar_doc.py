from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from funcao import gerar_atributos


def criar_doc(texto):
    lista_tabelas,lista_select,lista_from,lista_joins,lista_where,lista_group,lista_having,lista_order=gerar_atributos(texto=texto)



    document = Document()

    document.add_heading('Documentação SQL', 0)

    p = document.add_paragraph('este documento foi gerado de maneira automatica, por tanto recomendamos que observe se todas as tabelas se encontram.')

    document.add_heading('Tabelas:', 3)
    for x in lista_tabelas:
        document.add_paragraph(x, style='List Bullet')


    document.add_page_break()

    for x in range(0,len(lista_select)):
        document.add_heading(lista_tabelas[x], level=1)
        document.add_paragraph('Princial Origem: {}'.format(lista_from[x][0]), style='Intense Quote')

        document.add_paragraph('Informações da tabela:', style='List Bullet')

        def bonito(table):
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(10)

            row = table.rows[0]
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(15)




            document.add_paragraph('')    


        def gerar_tabela_origens(lista, Nome_atr1,Nome_atr2,Nome_atr3):
            if len(lista) != 0:
                lista.reverse()
                table = document.add_table(rows=1, cols=3)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = Nome_atr1
                hdr_cells[1].text = Nome_atr2
                hdr_cells[2].text = Nome_atr3
                for a, b, c in lista:
                    row_cells = table.add_row().cells
                    row_cells[0].text = a
                    row_cells[1].text = b
                    row_cells[2].text = c
                for cell in table.columns[0].cells:
                    cell.width = Inches(2)
                for cell in table.columns[2].cells:
                    cell.width = Inches(4)
                bonito(table)

        def gerar_tabela(lista, Nome_atr):
            if len(lista) != 0:
                lista = [item.rstrip().lstrip() for item in lista]
                table = document.add_table(rows=1, cols=1)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = Nome_atr
                for refs in lista:
                    row_cells = table.add_row().cells
                    row_cells[0].text = refs
                
                bonito(table)
        


                




        gerar_tabela(lista_select[x], 'Colunas:')
        gerar_tabela_origens(lista_joins[x],'Nome da Tabela','Apelido','Junção')
        gerar_tabela(lista_where[x], 'Filtro:')
        gerar_tabela(lista_group[x], 'Agrupamento:')
        gerar_tabela(lista_having[x], 'Filtro Agrupado:')
        gerar_tabela(lista_order[x], 'Ordenação:')
        
        
        document.add_page_break()


    document.save('Documentacao.docx')